import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
import threading
import time
import os
import pandas as pd
from datetime import datetime, date, timedelta
import ctypes
import matplotlib.pyplot as plt
from openpyxl import Workbook
from openpyxl.drawing.image import Image as OpenpyxlImage
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Font
import io
import zipfile
import pyautogui

# --- Optional Imports Check ---
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx2pdf import convert
    from pypdf import PdfWriter, PdfReader
except ImportError:
    # We will handle missing libraries gracefully in the UI if needed
    pass

# ==========================================
# PART 1: WORK LOG LOGIC (Refactored)
# ==========================================
class WorkLogTab:
    def __init__(self, parent_frame, username, data_dir):
        self.master = parent_frame
        self.username = username
        self.user_data_dir = data_dir
        
        self.is_monitoring = False
        self.is_paused = False
        self.log_data = []
        
        self.activity_excel_path = None
        self.summary_excel_path = None

        self.setup_ui()

    def setup_ui(self):
        # Header
        self.time_label = tk.Label(self.master, text="", font=("Arial", 12))
        self.time_label.pack(pady=10)
        self.update_time()

        self.status_label = tk.Label(self.master, text="Status: Ready", font=("Arial", 14, "bold"))
        self.status_label.pack(pady=10)

        # Buttons
        btn_frame = tk.Frame(self.master)
        btn_frame.pack(pady=10)

        self.clock_in_btn = tk.Button(btn_frame, text="Clock-in", font=("Arial", 11, "bold"), width=15, bg="#4CAF50", fg="white", command=self.clock_in)
        self.clock_in_btn.grid(row=0, column=0, padx=5, pady=5)
        
        self.break_btn = tk.Button(btn_frame, text="Take a Break", font=("Arial", 11, "bold"), width=15, bg="#FFC107", fg="black", command=self.take_a_break)
        self.break_btn.grid(row=1, column=0, padx=5, pady=5)

        self.clock_out_btn = tk.Button(btn_frame, text="Clock-out", font=("Arial", 11, "bold"), width=15, bg="#F44336", fg="white", command=self.clock_out)
        self.clock_out_btn.grid(row=2, column=0, padx=5, pady=5)

    def update_time(self):
        try:
            now = datetime.now()
            date_str = now.strftime("%A, %B %d, %Y")
            time_str = now.strftime("%I:%M:%S %p")
            self.time_label.config(text=f"User: {self.username}\n{date_str}\n{time_str}")
            self.master.after(1000, self.update_time)
        except:
            pass

    def clock_in(self):
        if not self.is_monitoring:
            self.is_monitoring = True
            self.is_paused = False
            self.start_time = datetime.now()
            self.log_activity("Clocked in to work", self.start_time, self.start_time)
            self.status_label.config(text="Status: Working...")
            self.monitoring_thread = threading.Thread(target=self.monitor_activity, daemon=True)
            self.monitoring_thread.start()
            messagebox.showinfo("Clocked In", "You have successfully clocked in.")
        elif self.is_paused:
            break_end_time = datetime.now()
            break_duration = break_end_time - self.break_start_time
            self.log_activity("Took a break from work", self.break_start_time, break_end_time, break_duration)
            self.is_paused = False
            self.status_label.config(text="Status: Working...")
            messagebox.showinfo("Resuming Work", "Your break has ended. Resuming work now.")
        else:
            messagebox.showinfo("Info", "You are already clocked in.")

    def take_a_break(self):
        if not self.is_monitoring or self.is_paused:
            messagebox.showwarning("Warning", "You must be clocked in to take a break.")
            return

        end_time = datetime.now()
        duration = end_time - self.start_time
        self.log_activity("Worked", self.start_time, end_time, duration)
        
        self.is_paused = True
        self.break_start_time = datetime.now()
        self.status_label.config(text="Status: On Break")
        messagebox.showinfo("On Break", "Enjoy your break! Click 'Clock-in' to resume work.")
    
    def clock_out(self):
        if not self.is_monitoring:
            messagebox.showwarning("Warning", "You must be clocked in to clock out.")
            return

        if self.is_paused:
            duration = datetime.now() - self.break_start_time
            self.log_activity("Took a break from work", self.break_start_time, datetime.now(), duration)
        else:
            duration = datetime.now() - self.start_time
            self.log_activity("Worked", self.start_time, datetime.now(), duration)

        self.is_monitoring = False
        self.is_paused = False
        self.log_activity("Clocked out from work", datetime.now(), datetime.now())
        
        self.save_data_to_excel()
        self.save_summary_to_excel()
        
        if self.activity_excel_path and self.summary_excel_path:
            self.create_zip_file()
            
        self.status_label.config(text="Status: Clocked Out")
        messagebox.showinfo("Clocked Out", f"Reports saved to Downloads folder.")

    # ... [Helper methods for WorkLogTab kept compact for readability] ...
    def log_activity(self, activity_type, start_t, end_t, duration=None):
        if duration:
            total_seconds = int(duration.total_seconds())
            hours, remainder = divmod(total_seconds, 3600)
            minutes, seconds = divmod(remainder, 60)
            duration_str = f"{activity_type} for {hours:02d}h {minutes:02d}m {seconds:02d}s"
        else:
            duration_str = activity_type
        
        self.log_data.append({
            "Employee": self.username, "Date": start_t.strftime("%d %B %Y"),
            "Start Time": start_t.strftime("%I:%M %p"), "End Time": end_t.strftime("%I:%M %p") if end_t else "-",
            "Activity Duration": duration_str
        })

    def is_internet_connected(self):
        try:
            import socket
            socket.create_connection(("8.8.8.8", 53), timeout=3)
            return True
        except: return False

    def is_system_locked(self):
        try:
            user32 = ctypes.windll.User32
            return user32.GetForegroundWindow() == 0
        except: return False
            
    def monitor_activity(self):
        last_internet = self.is_internet_connected()
        last_locked = self.is_system_locked()
        net_start = None
        lock_start = None
        
        while self.is_monitoring and not self.is_paused:
            time.sleep(1)
            curr_net = self.is_internet_connected()
            if not curr_net and last_internet: net_start = datetime.now()
            elif curr_net and not last_internet and net_start:
                self.log_activity("Internet interrupted", net_start, datetime.now(), datetime.now() - net_start)
                net_start = None
            last_internet = curr_net

            curr_locked = self.is_system_locked()
            if curr_locked and not last_locked: lock_start = datetime.now()
            elif not curr_locked and last_locked and lock_start:
                self.log_activity("System locked", lock_start, datetime.now(), datetime.now() - lock_start)
                lock_start = None
            last_locked = curr_locked

    def parse_duration(self, d_str):
        try:
            if " for " in d_str:
                parts = d_str.split(" for ")[1].strip().split(' ')
                return timedelta(hours=int(parts[0][:-1]), minutes=int(parts[1][:-1]), seconds=int(parts[2][:-1]))
        except: pass
        return timedelta(0)

    def analyze_data(self, df):
        # Simplified analysis for brevity in merged code
        df['DurSec'] = df['Activity Duration'].apply(lambda x: self.parse_duration(x).total_seconds())
        work_sec = df[df['Activity Duration'].str.contains("Worked for")]['DurSec'].sum()
        break_sec = df[df['Activity Duration'].str.contains("Took a break")]['DurSec'].sum()
        return {
            'Work Str': str(timedelta(seconds=int(work_sec))),
            'Break Count': df[df['Activity Duration'].str.contains("Took a break")].shape[0],
            'Work Hours': work_sec/3600, 'Break Hours': break_sec/3600
        }

    def save_data_to_excel(self):
        if not self.log_data: return
        try:
            df = pd.DataFrame(self.log_data)
            fname = f"{self.username.replace(' ', '')}-Work-Log-{date.today()}.xlsx"
            self.activity_excel_path = os.path.join(os.path.expanduser("~"), "Downloads", fname)
            df.to_excel(self.activity_excel_path, index=False) # Simplified for merged version
        except Exception as e: messagebox.showerror("Error", str(e))

    def save_summary_to_excel(self):
        if not self.log_data: return
        try:
            df = pd.DataFrame(self.log_data)
            analysis = self.analyze_data(df)
            # Create simple summary
            wb = Workbook()
            ws = wb.active
            ws.title = "Summary"
            ws['A1'] = f"Report for {self.username}"
            ws['A2'] = f"Total Work: {analysis['Work Str']}"
            
            fname = f"{self.username.replace(' ', '')}-Summary-{date.today()}.xlsx"
            self.summary_excel_path = os.path.join(os.path.expanduser("~"), "Downloads", fname)
            wb.save(self.summary_excel_path)
        except Exception as e: messagebox.showerror("Error", str(e))

    def create_zip_file(self):
        try:
            zname = f"{self.username.replace(' ', '')}-Reports-{date.today()}.zip"
            zpath = os.path.join(os.path.expanduser("~"), "Downloads", zname)
            with zipfile.ZipFile(zpath, 'w', zipfile.ZIP_DEFLATED) as zf:
                if self.activity_excel_path: zf.write(self.activity_excel_path, os.path.basename(self.activity_excel_path))
                if self.summary_excel_path: zf.write(self.summary_excel_path, os.path.basename(self.summary_excel_path))
            if self.activity_excel_path: os.remove(self.activity_excel_path)
            if self.summary_excel_path: os.remove(self.summary_excel_path)
        except: pass

# ==========================================
# PART 2: SCREENSHOT LOGIC (Refactored)
# ==========================================
class ScreenshotTab:
    def __init__(self, parent_frame, username, data_dir):
        self.master = parent_frame
        self.username = username
        self.status = "stopped"
        self.thread = None
        self.report_password = "MaCroDmT" # Hardcoded password
        
        self.screenshot_dir = os.path.join(data_dir, "screenshots")
        os.makedirs(self.screenshot_dir, exist_ok=True)
        
        self.setup_ui()

    def setup_ui(self):
        # Greeting
        self.greeting_label = tk.Label(self.master, text=f"Monitoring: {self.username}", font=("Arial", 12, "bold"))
        self.greeting_label.pack(pady=10)

        # Controls
        ctrl_frame = tk.Frame(self.master)
        ctrl_frame.pack(pady=10)

        self.start_btn = tk.Button(ctrl_frame, text="Start Capture", width=15, command=self.start, bg="#4CAF50", fg="white", font=("Arial", 10, "bold"))
        self.start_btn.grid(row=0, column=0, pady=5)

        self.pause_btn = tk.Button(ctrl_frame, text="Pause", width=15, command=self.pause, bg="#FFEB3B", fg="black", font=("Arial", 10, "bold"))
        self.pause_btn.grid(row=1, column=0, pady=5)

        self.stop_btn = tk.Button(ctrl_frame, text="Stop & Report", width=15, command=self.stop, bg="#F44336", fg="white", font=("Arial", 10, "bold"))
        self.stop_btn.grid(row=2, column=0, pady=5)
        
        lbl_note = tk.Label(self.master, text="Screenshots taken every 5 mins", font=("Arial", 8, "italic"))
        lbl_note.pack(pady=5)

    def capture_screenshots(self):
        while self.status == "started":
            now = datetime.now()
            filename = f"screenshot_{now.strftime('%Y%m%d_%H%M%S')}.png"
            filepath = os.path.join(self.screenshot_dir, filename)
            try:
                pyautogui.screenshot().save(filepath)
            except Exception as e: print(f"SS Error: {e}")

            # Wait 5 mins (300 secs) checking status
            for _ in range(300): 
                if self.status != "started": break
                time.sleep(1)

    def start(self):
        if self.status != "started":
            self.status = "started"
            if self.thread is None or not self.thread.is_alive():
                self.thread = threading.Thread(target=self.capture_screenshots, daemon=True)
                self.thread.start()
            messagebox.showinfo("SS Status", "Screenshot capture started.")

    def pause(self):
        if self.status == "started":
            self.status = "paused"
            messagebox.showinfo("SS Status", "Screenshot capture paused.")

    def stop(self):
        if self.status != "stopped":
            self.status = "stopped"
            if self.thread and self.thread.is_alive():
                self.thread.join(timeout=1.0)
            
            self.create_report()

    def create_report(self):
        # Dependencies check
        if 'Document' not in globals() or 'convert' not in globals():
            messagebox.showerror("Error", "MS Word/docx libraries not installed/found.")
            return

        today_str = date.today().strftime("%Y%m%d")
        files = [os.path.join(self.screenshot_dir, f) for f in os.listdir(self.screenshot_dir) 
                 if f.startswith(f"screenshot_{today_str}") and f.endswith(".png")]
        
        if not files:
            messagebox.showinfo("Info", "No screenshots to save.")
            return

        files.sort()
        doc = Document()
        
        # Add images 2 per page
        for i in range(0, len(files), 2):
            if i > 0: doc.add_page_break()
            try:
                doc.add_paragraph().add_run().add_picture(files[i], width=Inches(6.0))
                if i + 1 < len(files):
                    doc.add_paragraph().add_run().add_picture(files[i+1], width=Inches(6.0))
            except: pass

        base_name = f"screenshots_{self.username}_{today_str}"
        temp_doc = os.path.join(self.screenshot_dir, f"{base_name}_temp.docx")
        temp_pdf = os.path.join(self.screenshot_dir, f"{base_name}_temp.pdf")
        final_pdf = os.path.join(os.path.expanduser("~"), "Downloads", f"{base_name}_SECURE.pdf")

        try:
            doc.save(temp_doc)
            convert(temp_doc, temp_pdf)
            
            reader = PdfReader(temp_pdf)
            writer = PdfWriter()
            for page in reader.pages: writer.add_page(page)
            writer.encrypt(self.report_password)
            
            with open(final_pdf, "wb") as f: writer.write(f)
            
            # Cleanup
            os.remove(temp_doc)
            os.remove(temp_pdf)
            for f in files: os.remove(f)
            
            messagebox.showinfo("Success", f"Encrypted PDF saved to:\n{final_pdf}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to create PDF: {e}")

# ==========================================
# PART 3: MAIN APP CONTAINER
# ==========================================
class MainApplication:
    def __init__(self, root):
        self.root = root
        self.root.title("Employee Monitoring Suite")
        self.root.geometry("450x450")
        
        # Unified Data Directory
        self.data_dir = os.path.join(os.path.expanduser("~"), "WorkLogData")
        os.makedirs(self.data_dir, exist_ok=True)
        self.config_file = os.path.join(self.data_dir, "config.txt")
        
        self.username = self.load_username()
        
        if not self.username:
            self.show_login_frame()
        else:
            self.show_main_interface()

    def load_username(self):
        if os.path.exists(self.config_file):
            with open(self.config_file, 'r') as f:
                return f.read().strip()
        return None

    def save_username(self, name):
        with open(self.config_file, 'w') as f:
            f.write(name)
        self.username = name

    def show_login_frame(self):
        self.login_frame = tk.Frame(self.root)
        self.login_frame.pack(expand=True, fill='both', padx=20, pady=20)
        
        tk.Label(self.login_frame, text="Welcome Employee", font=("Arial", 16, "bold")).pack(pady=20)
        tk.Label(self.login_frame, text="Please enter your name:", font=("Arial", 12)).pack()
        
        self.name_entry = tk.Entry(self.login_frame, font=("Arial", 12))
        self.name_entry.pack(pady=10)
        
        tk.Button(self.login_frame, text="Start Session", command=self.submit_login, 
                  bg="#2196F3", fg="white", font=("Arial", 12, "bold")).pack(pady=20)

    def submit_login(self):
        name = self.name_entry.get().strip()
        if name:
            self.save_username(name)
            self.login_frame.destroy()
            self.show_main_interface()
        else:
            messagebox.showwarning("Error", "Name is required")

    def show_main_interface(self):
        # Create Tab Container
        notebook = ttk.Notebook(self.root)
        notebook.pack(expand=True, fill='both')

        # Tab 1: Work Log
        tab1 = tk.Frame(notebook, bg="#f0f0f0")
        notebook.add(tab1, text="  Work Log  ")
        self.work_log_app = WorkLogTab(tab1, self.username, self.data_dir)

        # Tab 2: Screenshots
        tab2 = tk.Frame(notebook)
        notebook.add(tab2, text="  Screen Monitor  ")
        self.screenshot_app = ScreenshotTab(tab2, self.username, self.data_dir)

if __name__ == "__main__":
    root = tk.Tk()
    # Optional: Handle High DPI displays
    try: ctypes.windll.shcore.SetProcessDpiAwareness(1)
    except: pass
    
    app = MainApplication(root)
    root.mainloop()